from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path


@dataclass(frozen=True)
class LoadedDocument:
    path: Path
    text: str
    kind: str


TEXT_EXTENSIONS = {
    ".txt": "text",
    ".md": "markdown",
    ".markdown": "markdown",
    ".rst": "text",
    ".csv": "text",
    ".json": "text",
}


def load_document(path: str | Path) -> LoadedDocument:
    document_path = Path(path)
    if not document_path.exists():
        raise FileNotFoundError(document_path)

    suffix = document_path.suffix.lower()
    if suffix in TEXT_EXTENSIONS:
        return LoadedDocument(
            path=document_path,
            text=document_path.read_text(encoding="utf-8"),
            kind=TEXT_EXTENSIONS[suffix],
        )
    if suffix == ".pdf":
        return LoadedDocument(path=document_path, text=_read_pdf(document_path), kind="pdf")
    if suffix == ".docx":
        return LoadedDocument(path=document_path, text=_read_docx(document_path), kind="docx")

    raise ValueError(f"Unsupported document extension: {suffix}")


def load_text_bundle(paths: list[str | Path], max_chars: int = 20000) -> str:
    parts: list[str] = []
    remaining = max_chars
    for path in paths:
        if remaining <= 0:
            break
        document = load_document(path)
        text = document.text.strip()
        if len(text) > remaining:
            text = text[:remaining] + "\n[truncated]"
        parts.append(f"## Source: {document.path}\n{text}")
        remaining -= len(text)
    return "\n\n".join(parts)


def _read_pdf(path: Path) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise RuntimeError("Install pypdf to read PDF files: pip install pypdf") from exc

    reader = PdfReader(str(path))
    pages: list[str] = []
    for index, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        pages.append(f"[page {index}]\n{text.strip()}")
    return "\n\n".join(pages)


def _read_docx(path: Path) -> str:
    try:
        from docx import Document as DocxDocument
    except ImportError as exc:
        raise RuntimeError("Install python-docx to read DOCX files: pip install python-docx") from exc

    doc = DocxDocument(str(path))
    paragraphs = [paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()]
    table_rows: list[str] = []
    for table in doc.tables:
        for row in table.rows:
            table_rows.append(" | ".join(cell.text.strip() for cell in row.cells))
    return "\n\n".join(paragraphs + table_rows)
